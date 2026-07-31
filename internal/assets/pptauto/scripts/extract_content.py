"""
源内容提取脚本：从 PDF/DOCX/XLSX/CSV/TXT 提取文本内容。

用法: python extract_content.py <input_file> [input_file2 ...] <output_json>

支持格式：
- .pdf → 提取文字（需要 PyMuPDF 或 pdfminer）
- .docx/.doc → 提取文字和表格
- .xlsx/.xls → 提取表格数据
- .csv → 提取表格数据
- .txt/.md → 直接读取（自动检测编码）

输出 JSON:
{
  "sources": ["文件路径"],
  "formats": ["pdf|docx|xlsx|csv|txt"],
  "title": "文档标题",
  "content": "合并的文本内容",
  "sections": [{"title": "章节标题", "content": "章节内容"}],
  "tables": [{"headers": [...], "rows": [[...]]}],
  "files": [{"source": "...", "format": "...", "title": "...", ...}]
}
"""
import json, sys, os, re, csv


def detect_encoding(path):
    """检测文件编码"""
    try:
        import chardet
        with open(path, "rb") as f:
            raw = f.read(10000)
        result = chardet.detect(raw)
        return result.get("encoding", "utf-8") or "utf-8"
    except:
        return "utf-8"


def extract_pdf(path):
    """从 PDF 提取文字"""
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text, []
    except ImportError:
        pass

    try:
        from pdfminer.high_level import extract_text
        return extract_text(path), []
    except ImportError:
        pass

    return None, []


def extract_docx(path):
    """从 DOCX 提取文字和表格"""
    text = ""
    tables = []

    try:
        from docx import Document
        doc = Document(path)

        # 提取段落
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"

        # 提取表格
        for table in doc.tables:
            headers = []
            rows = []
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                if i == 0:
                    headers = cells
                else:
                    rows.append(cells)
            if headers:
                tables.append({"headers": headers, "rows": rows})

        return text, tables
    except ImportError:
        pass

    # 兜底：zipfile
    try:
        import zipfile
        from xml.etree import ElementTree as ET
        with zipfile.ZipFile(path) as z:
            with z.open("word/document.xml") as f:
                tree = ET.parse(f)
                root = tree.getroot()
                ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
                for para in root.iter(f"{{{ns['w']}}}p"):
                    for run in para.iter(f"{{{ns['w']}}}t"):
                        if run.text:
                            text += run.text
                    text += "\n"
        return text, []
    except:
        pass

    return None, []


def extract_xlsx(path):
    """从 XLSX 提取表格数据"""
    tables = []
    text = ""

    try:
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            # 第一行作为表头
            headers = [str(cell) if cell is not None else "" for cell in rows[0]]
            data_rows = []
            for row in rows[1:]:
                data_rows.append([str(cell) if cell is not None else "" for cell in row])
            tables.append({"headers": headers, "rows": data_rows, "sheet": sheet_name})
            # 也生成文本版本
            text += f"\n## {sheet_name}\n"
            for row in rows[:50]:  # 最多50行
                text += " | ".join(str(cell) if cell is not None else "" for cell in row) + "\n"
        wb.close()
        return text, tables
    except ImportError:
        pass

    return None, []


def extract_csv(path):
    """从 CSV 提取表格数据"""
    tables = []
    text = ""

    encoding = detect_encoding(path)
    try:
        with open(path, "r", encoding=encoding, errors="ignore") as f:
            reader = csv.reader(f)
            rows = list(reader)
            if rows:
                headers = rows[0]
                data_rows = rows[1:]
                tables.append({"headers": headers, "rows": data_rows})
                for row in rows[:50]:
                    text += " | ".join(row) + "\n"
        return text, tables
    except:
        pass

    return None, []


def extract_txt(path):
    """读取纯文本"""
    encoding = detect_encoding(path)
    with open(path, "r", encoding=encoding, errors="ignore") as f:
        return f.read(), []


def extract_sections(text):
    """将文本拆分为章节"""
    lines = text.split("\n")
    sections = []
    current_title = ""
    current_content = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        is_title = False
        if re.match(r'^[一二三四五六七八九十]+[、.．]', line):
            is_title = True
        elif re.match(r'^第[一二三四五六七八九十]+[章节部分]', line):
            is_title = True
        elif re.match(r'^\d+[、.．]\s*', line):
            is_title = True
        elif re.match(r'^[（(]\d+[）)]\s*', line):
            is_title = True
        elif len(line) < 30 and line == line.upper() and len(line) > 2:
            is_title = True
        elif re.match(r'^#{1,3}\s', line):
            is_title = True

        if is_title:
            if current_title or current_content:
                sections.append({
                    "title": current_title,
                    "content": "\n".join(current_content)
                })
            current_title = line.lstrip("#").strip()
            current_content = []
        else:
            current_content.append(line)

    if current_title or current_content:
        sections.append({
            "title": current_title,
            "content": "\n".join(current_content)
        })

    return sections


def extract_single(path):
    """提取单个文件的内容"""
    ext = os.path.splitext(path)[1].lower()

    text = None
    tables = []
    fmt = None

    if ext == ".pdf":
        text, tables = extract_pdf(path)
        fmt = "pdf"
    elif ext in (".docx", ".doc"):
        text, tables = extract_docx(path)
        fmt = "docx"
    elif ext in (".xlsx", ".xls"):
        text, tables = extract_xlsx(path)
        fmt = "xlsx"
    elif ext == ".csv":
        text, tables = extract_csv(path)
        fmt = "csv"
    elif ext in (".txt", ".md", ".markdown"):
        text, tables = extract_txt(path)
        fmt = "txt"
    else:
        return None

    if not text and not tables:
        return None

    title = os.path.splitext(os.path.basename(path))[0]
    lines = [l.strip() for l in (text or "").split("\n") if l.strip()]
    if lines and len(lines[0]) < 50:
        title = lines[0]

    sections = extract_sections(text or "")

    return {
        "source": path,
        "format": fmt,
        "title": title,
        "content": (text or "")[:5000],
        "sections": sections[:20],
        "tables": tables[:10],
    }


def main():
    if len(sys.argv) < 3:
        print("用法: python extract_content.py <file1> [file2 ...] <output.json>")
        print("支持: PDF, DOCX, XLSX, CSV, TXT, MD")
        sys.exit(1)

    output_path = sys.argv[-1]
    input_paths = sys.argv[1:-1]

    if not input_paths:
        print("ERROR: 未指定输入文件")
        sys.exit(1)

    results = []
    combined_content = ""
    all_tables = []

    for input_path in input_paths:
        if not os.path.exists(input_path):
            print(f"  跳过: {input_path} (不存在)")
            continue

        print(f"  提取: {os.path.basename(input_path)}...", end=" ")
        result = extract_single(input_path)
        if result:
            results.append(result)
            combined_content += f"\n\n## {result['title']}\n\n{result['content']}"
            all_tables.extend(result.get("tables", []))
            print(f"OK ({result['format']}, {len(result['content'])} 字符, {len(result.get('tables',[]))} 表格)")
        else:
            print("FAILED")

    if not results:
        print("ERROR: 无法提取任何内容")
        sys.exit(1)

    output = {
        "sources": [r["source"] for r in results],
        "formats": [r["format"] for r in results],
        "title": results[0]["title"] if len(results) == 1 else f"{results[0]['title']}等{len(results)}份文档",
        "content": combined_content.strip()[:8000],
        "sections": [],
        "tables": all_tables,
        "files": results,
    }

    for r in results:
        for s in r.get("sections", []):
            output["sections"].append(s)

    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"\n  共提取 {len(results)} 个文件")
    print(f"  总文字: {len(combined_content)} 字符")
    print(f"  章节: {len(output['sections'])} 个")
    print(f"  表格: {len(all_tables)} 个")
    print(f"  输出: {output_path}")


if __name__ == "__main__":
    main()
